import os
import sys
import shutil
import sqlite3
import logging
import platform
import psutil
import traceback
import re
from pathlib import Path
from docx import Document
import arcpy
from datetime import datetime
import sys
sys.stdout.reconfigure(encoding='utf-8')
class WordAutomation:
    def __init__(self):
        """Inicializa la automatización con mejor manejo de archivos temporales"""
        try:
            # Auto-detectar la ruta raíz
            self.root_path = self._find_root_path()
            
            # Configurar rutas principales
            self.temp_root = self.root_path / "Files" / "Temporary_Files" / "MODELO_IGAC"
            self.process_temp_path = self.temp_root / "word_automation_temp"
            
            # Configurar logging en Temporary_Files
            log_file = self.temp_root / "word_automation_urbano.log"
            logging.basicConfig(
                level=logging.INFO,
                format='%(asctime)s - %(levelname)s - %(message)s',
                handlers=[
                    logging.FileHandler(log_file),
                    logging.StreamHandler()
                ]
            )
            
            # Crear carpeta temporal para el proceso si no existe
            self.process_temp_path.mkdir(exist_ok=True)
            
            print(f"\n{'='*50}")
            print("Configuración inicial:")
            print(f"Ruta raíz detectada: {self.root_path}")
            
            # Datasets a procesar
            self.datasets_to_process = [
                #"RURAL",
                "URBANO"
            ]
            print(f"Datasets a procesar: {', '.join(self.datasets_to_process)}")
            
            # Configurar rutas de bases de datos
            self.db_paths = {
                'conteo': self.temp_root / "db" / "conteo_elementos.db",
                'errores': self.temp_root / "db" / "errores_consistencia_formato.db",
                'excepciones': self.temp_root / "db" / "excepciones_consistencia_formato.db",
                'omision': self.temp_root / "db" / "omision_comision.db",
                'municipios': self.root_path / "Files" / "Municipios" / "municipios.db",
                'registro_errores': self.temp_root / "db" / "registro_errores.db"
            }
            
            # Verificar existencia de bases de datos
            print("\nVerificando rutas de bases de datos:")
            missing_dbs = []
            for db_name, db_path in self.db_paths.items():
                exists = db_path.exists()
                print(f"{db_name}: {'[OK]' if exists else '[X]'} {db_path}")
                if not exists:
                    missing_dbs.append(db_name)
            
            if missing_dbs:
                raise FileNotFoundError(f"Bases de datos faltantes: {', '.join(missing_dbs)}")
            
            print(f"\n{'='*50}")
            
        except Exception as e:
            print(f"\n[ERROR CRÍTICO] Error durante la inicialización: {e}")
            raise

    def _find_root_path(self):
        """
        Encuentra la ruta raíz del proyecto verificando la estructura de directorios esperada.
        """
        try:
            current = Path(os.path.abspath(__file__)).parent
            
            while current.parent != current:
                # Verifica la estructura característica del proyecto
                required_paths = [
                    current / "Files" / "Temporary_Files" / "MODELO_IGAC",
                    current / "Files" / "Temporary_Files" / "array_config.txt",
                    current / "Files" / "Templates" /"MODELO_IGAC"/ "04_REPORTE_FINAL",
                    current / "Files" / "Municipios" / "municipios.db"
                ]
                
                if all(path.exists() for path in required_paths):
                    print(f"Raíz del proyecto encontrada en: {current}")
                    return current
                
                current = current.parent
            
            raise Exception(
                "No se encontró la raíz del proyecto.\n"
                "Verifique que está ejecutando el script desde dentro del proyecto "
                "y que existe la siguiente estructura:\n"
                "- Files/Temporary_Files/MODELO_IGAC\n"
                "- Files/Temporary_Files/array_config.txt\n"
                "- Files/Templates/MODELO_IGAC/04_REPORTE_FINAL\n"
                "- Files/Municipios/municipios.db"
            )
        except Exception as e:
            print(f"[ERROR] Error encontrando ruta raíz: {e}")
            raise

    def verify_paths(self):
        """Verifica que todas las rutas necesarias existan"""
        try:
            print("\nVerificando estructura de carpetas:")
            paths_to_check = [
                self.root_path / "Files" / "Templates" / "MODELO_IGAC"/ "04_REPORTE_FINAL",
                self.temp_root,
                self.temp_root / "db"
            ]
            
            missing_paths = []
            for path in paths_to_check:
                exists = path.exists()
                print(f"{'[OK]' if exists else '[X]'} {path}")
                if not exists:
                    missing_paths.append(str(path))
            
            if missing_paths:
                print("\n[ADVERTENCIA] Las siguientes rutas no existen:")
                for path in missing_paths:
                    print(f"  - {path}")
                return False
            
            return True
            
        except Exception as e:
            print(f"[ERROR] Error verificando rutas: {e}")
            return False

    def copy_templates(self):
        """Copia las plantillas a la carpeta temporal del proceso"""
        try:
            print("\nCopiando plantillas:")
            
            for dataset in self.datasets_to_process:
                template_path = self.root_path / "Files" / "Templates" / "MODELO_IGAC"/ "04_REPORTE_FINAL" / dataset
                
                print(f"\nProcesando dataset {dataset}:")
                print(f"Buscando plantillas en: {template_path}")
                
                if not template_path.exists():
                    print(f"  [ADVERTENCIA] Ruta de plantilla no encontrada: {template_path}")
                    continue
                
                docx_files = list(template_path.glob("*.docx"))
                if not docx_files:
                    print(f"  [ADVERTENCIA] No se encontraron archivos .docx en {template_path}")
                    continue
                    
                for docx_file in docx_files:
                    try:
                        # Copiar a nuestra carpeta temporal específica
                        shutil.copy2(docx_file, self.process_temp_path)
                        print(f"  [OK] Copiado: {docx_file.name} -> {self.process_temp_path}")
                    except Exception as e:
                        print(f"  [ERROR] Error copiando {docx_file.name}: {e}")
                        
        except Exception as e:
            print(f"[ERROR] Error copiando plantillas: {e}")
            raise

    def execute_sql_query(self, db_path, query, params=None):
        """Ejecuta una consulta SQL y retorna el resultado con mejor manejo de errores"""
        try:
            if not Path(db_path).exists():
                raise FileNotFoundError(f"Base de datos no encontrada: {db_path}")
                
            with sqlite3.connect(db_path) as conn:
                cursor = conn.cursor()
                if params:
                    result = cursor.execute(query, params).fetchone()
                else:
                    result = cursor.execute(query).fetchone()
                    
                if result is None:
                    print(f"[ADVERTENCIA] La consulta no retornó resultados: {query}")
                    return 0
                    
                return result[0] if result else 0
                
        except sqlite3.Error as e:
            print(f"[ERROR] Error en la base de datos: {e}")
            print(f"Base de datos: {db_path}")
            print(f"Query: {query}")
            print(f"Parámetros: {params}")
            return 0
        except Exception as e:
            print(f"[ERROR] Error ejecutando consulta SQL: {e}")
            return 0

    def get_gdb_name(self):
        """Obtiene el nombre de la geodatabase en Temporary_Files con validación mejorada"""
        try:
            print("\nBuscando geodatabase:")
            print(f"Buscando en: {self.temp_root}")
            
            if not self.temp_root.exists():
                raise FileNotFoundError(f"Ruta temporal no encontrada: {self.temp_root}")
            
            gdbs = list(self.temp_root.glob("*.gdb"))
            
            if not gdbs:
                print("  [X] No se encontró ninguna geodatabase")
                return None
                
            if len(gdbs) > 1:
                print(f"  [ADVERTENCIA] Se encontraron múltiples geodatabases: {[gdb.name for gdb in gdbs]}")
                print(f"  Se utilizará la primera: {gdbs[0].name}")
                
            print(f"  [OK] Geodatabase encontrada: {gdbs[0].name}")
            return gdbs[0].name
            
        except Exception as e:
            print(f"[ERROR] Error buscando geodatabase: {e}")
            return None
        
    def get_municipality_info(self, dataset):
        """Obtiene información del municipio y departamento desde la GDB con validación mejorada"""
        try:
            gdb_name = self.get_gdb_name()
            if not gdb_name:
                raise ValueError("No se pudo obtener el nombre de la geodatabase")

            print(f"\nObteniendo información del municipio para dataset {dataset}")
            gdb_path = self.temp_root / gdb_name
            
            if not gdb_path.exists():
                raise FileNotFoundError(f"Geodatabase no encontrada: {gdb_path}")
            
            print(f"Conectando a GDB: {gdb_path}")
            arcpy.env.workspace = str(gdb_path)
            
            # Verificar existencia del dataset
            if not arcpy.Exists(dataset):
                raise ValueError(f"Dataset {dataset} no encontrado")
                
            print(f"[OK] Dataset encontrado: {dataset}")
            
            # Determinar el feature class a buscar
            feature_class = "U_MANZANA" if "URBANO" in dataset else "R_VEREDA"
            print(f"Buscando feature class: {feature_class}")
            
            # Lista de feature classes en el dataset
            feature_classes = arcpy.ListFeatureClasses(feature_dataset=dataset)
            print(f"Feature classes encontrados en el dataset: {feature_classes}")
            
            if feature_class not in feature_classes:
                raise ValueError(f"Feature class {feature_class} no encontrado")
            
            print(f"[OK] Feature class encontrado: {feature_class}")
            
            # Obtener códigos de municipio
            codes = set()
            with arcpy.da.SearchCursor(feature_class, ["CODIGO_MUNICIPIO"]) as cursor:
                for i, row in enumerate(cursor):
                    if i < 5 and row[0]:  # Tomar hasta 5 registros no nulos
                        codes.add(row[0])
                    if i >= 5:
                        break
            
            if not codes:
                raise ValueError("No se encontraron códigos de municipio válidos")

            codigo_municipio = codes.pop()
            print(f"Código de municipio encontrado: {codigo_municipio}")
            
            # Buscar en la base de datos de municipios
            with sqlite3.connect(self.db_paths['municipios']) as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT MUNICIPIO, DEPARTAMENTO 
                    FROM municipios 
                    WHERE CODIGO_DANE = ?
                """, (codigo_municipio,))
                result = cursor.fetchone()
                
                if not result:
                    raise ValueError(f"No se encontró información para el código de municipio: {codigo_municipio}")
                
                print(f"[OK] Municipio: {result[0]}, Departamento: {result[1]}")
                return result[0], result[1]
                
        except arcpy.ExecuteError as e:
            print(f"[ERROR] Error de ArcPy: {e}")
            print(arcpy.GetMessages(2))
            return None, None
        except Exception as e:
            print(f"[ERROR] Error obteniendo información del municipio: {e}")
            return None, None

    def get_conteo_variables(self, dataset, municipio):
        """Obtiene todas las variables de conteo para un dataset con validación mejorada"""
        print(f"\nObteniendo variables de conteo para {dataset}")
        
        variables = {}
        
        # Lista de variables base con sus nombres formateados
        base_vars = [
            'U_BARRIO',
            'U_CONSTRUCCION',
            
            'U_MANZANA',
            'U_NOMENCLATURA_DOMICILIARIA',
            'U_NOMENCLATURA_VIAL',
            'U_PERIMETRO',
            'U_SECTOR',
            'U_TERRENO',
            
            'U_UNIDAD',
            
            'U_ZONA_HOMOGENEA_GEOECONOMICA',
            'U_ZONA_HOMOGENEA_FISICA'
        ]
        
        try:
            # Procesar cada variable base y su versión con guion bajo
            print("\nProcesando variables individuales:")
            for var_name in base_vars:
                try:
                    # Consulta para obtener el valor de la variable
                    query = f"SELECT {var_name} FROM {dataset} LIMIT 1"
                    value = self.execute_sql_query(self.db_paths['conteo'], query)
                    variables[var_name] = value
                    
                    # Procesar versión con guion bajo
                    var_underscore = f"{var_name}_"
                    variables[var_underscore] = "" if value > 0 else "Vacía"
                    
                    print(f"[OK] {var_name}: {value} -> {var_underscore}: '{variables[var_underscore]}'")
                    
                except Exception as e:
                    print(f"[ERROR] Error procesando {var_name}: {e}")
                    variables[var_name] = 0
                    variables[f"{var_name}_"] = "Vacía"

            # Calcular total_conteo
            try:
                total_query = f"SELECT ({' + '.join(base_vars)}) FROM {dataset} LIMIT 1"
                variables['total_conteo'] = self.execute_sql_query(self.db_paths['conteo'], total_query)
                print(f"\nTotal conteo calculado: {variables['total_conteo']}")
                
            except Exception as e:
                print(f"[ERROR] Error calculando total_conteo: {e}")
                # Si falla la consulta, sumar los valores individuales
                variables['total_conteo'] = sum(variables[var] for var in base_vars)
                print(f"Total conteo calculado por suma: {variables['total_conteo']}")
            
            return variables
            
        except Exception as e:
            print(f"[ERROR] Error general en get_conteo_variables: {e}")
            return variables
        
    def get_consistencia_variables(self, dataset):
        """Obtiene todas las variables relacionadas con consistencia con manejo mejorado para urbano"""
        print(f"\nObteniendo variables de consistencia para {dataset}")
        results = {}
        
        try:
            # Consultas para registro_errores.db adaptadas para urbano
            error_queries = {
                'no_huecos': """
                    SELECT COUNT(*) AS cantidad 
                    FROM URBANO 
                    WHERE RuleDescription IN ('Must Not Have Gaps', 'No debe tener espacios')
                """,
                'no_superponer': """
                    SELECT COUNT(*) AS cantidad 
                    FROM URBANO 
                    WHERE RuleDescription IN ('Must Not Overlap', 'No debe superponerse')
                """,
                'cubierto_por': """
                    SELECT COUNT(*) AS cantidad 
                    FROM URBANO 
                    WHERE RuleDescription IN ('Must Be Covered By Feature Class Of', 'Debe ser cubierto por la clase de entidad de')
                """,
                'cubrirse_entre': """
                    SELECT COUNT(*) AS cantidad 
                    FROM URBANO 
                    WHERE RuleDescription IN ('Must Cover Each Other', 'Deben cubrirse entre ellos')
                """
            }

            # Consultas para excepciones adaptadas para urbano
            exception_queries = {
                'no_huecos_': """
                    SELECT COUNT(*) 
                    FROM URBANO 
                    WHERE RuleDescription IN ('Must Not Have Gaps', 'No debe tener espacios')
                    AND IsException != 0
                """,
                'no_superponer_': """
                    SELECT COUNT(*) 
                    FROM URBANO 
                    WHERE RuleDescription IN ('Must Not Overlap', 'No debe superponerse')
                    AND IsException != 0
                """,
                'cubierto_por_': """
                    SELECT COUNT(*) 
                    FROM URBANO 
                    WHERE RuleDescription IN ('Must Be Covered By Feature Class Of', 'Debe ser cubierto por la clase de entidad de')
                    AND IsException != 0
                """,
                'cubrirse_entre_': """
                    SELECT COUNT(*) 
                    FROM URBANO 
                    WHERE RuleDescription IN ('Must Cover Each Other', 'Deben cubrirse entre ellos')
                    AND IsException != 0
                """
            }
            
            # Procesar errores desde registro_errores.db
            print("Procesando errores desde registro_errores.db...")
            for var_name, query in error_queries.items():
                try:
                    results[var_name] = self.execute_sql_query(self.db_paths['registro_errores'], query)
                    print(f"[OK] {var_name}: {results[var_name]}")
                except Exception as e:
                    print(f"[ERROR] Error procesando {var_name}: {e}")
                    results[var_name] = 0

            # Procesar excepciones
            print("\nProcesando excepciones...")
            for var_name, query in exception_queries.items():
                try:
                    results[var_name] = self.execute_sql_query(self.db_paths['registro_errores'], query)
                    print(f"[OK] {var_name}: {results[var_name]}")
                except Exception as e:
                    print(f"[ERROR] Error procesando {var_name}: {e}")
                    results[var_name] = 0
            
            

            # Calcular sumas con validación
            print("\nCalculando sumas...")
            try:
                results.update({
                    'sum1': results['no_huecos'] - results['no_huecos_'],
                    'sum2': results['no_superponer'] - results['no_superponer_'],
                    'sum3': results['cubierto_por'] - results['cubierto_por_'],
                    'sum4': results['cubrirse_entre'] - results['cubrirse_entre_'],
                    'sum_t_situa': sum(results[k] for k in ['no_huecos', 'no_superponer', 'cubierto_por', 'cubrirse_entre']),
                    'sum_cant_ex': sum(results[k] for k in ['no_huecos_', 'no_superponer_', 'cubierto_por_', 'cubrirse_entre_']),
                })
                
                results['sum_total_err'] = results['sum1'] + results['sum2'] + results['sum3'] + results['sum4']
                print(f"Sumas calculadas correctamente: sum_total_err = {results['sum_total_err']}")
            except Exception as e:
                print(f"[ERROR] Error calculando sumas: {e}")
                for key in ['sum1', 'sum2', 'sum3', 'sum4', 'sum_t_situa', 'sum_cant_ex', 'sum_total_err']:
                    results[key] = 0

            # Consulta para consistencia_total adaptada para urbano
            consistencia_query = """
            SELECT (
                U_UNIDAD + 
                U_TERRENO + 
                U_NOMENCLATURA_VIAL +
                U_TERRENO_U_MANZANA + 
                U_UNIDAD_U_TERRENO +
                U_NOMENCLATURA_DOMICILIARIA + 
                U_UNIDAD_U_CONSTRUCCION +
                U_NOMENCLATURA_DOMICILIARIA_U_TERRENO + 
                U_MANZANA +
                U_BARRIO + 
                U_CONSTRUCCION_U_TERRENO +
                U_CONSTRUCCION + 
                U_MANZANA_U_SECTOR +
                U_SECTOR
            ) AS total_primera_fila 
            FROM URBANO  
            LIMIT 1"""

            print("\nObteniendo consistencia total y excepciones...")
            try:
                results['consistencia_total'] = self.execute_sql_query(self.db_paths['errores'], consistencia_query)
                results['consistencia_except'] = self.execute_sql_query(self.db_paths['excepciones'], consistencia_query)
                
                # Calcular excepciones y total-errores
                results['consistencia_except'] = results['consistencia_total'] - results['consistencia_except']
                results['total-errores'] = results['consistencia_total'] - results['consistencia_except']
                
                print(f"consistencia_total: {results['consistencia_total']}")
                print(f"consistencia_except: {results['consistencia_except']}")
                print(f"total-errores: {results['total-errores']}")
            except Exception as e:
                print(f"[ERROR] Error calculando consistencia: {e}")
                results.update({
                    'consistencia_total': 0,
                    'consistencia_except': 0,
                    'total-errores': 0
                })
            
            # Añadir fecha
            results['date'] = datetime.now().strftime('%Y/%m/%d')
            
            return results
            
        except Exception as e:
            print(f"[ERROR] Error general en get_consistencia_variables: {e}")
            return results


        

    def process_dataset_variables(self, dataset):
        """Procesa todas las variables para un dataset urbano específico"""
        print(f"\nProcesando variables para dataset {dataset}")
        variables = {}
        
        try:
            # 1. Fecha actual
            variables['date'] = datetime.now().strftime('%Y/%m/%d')
            
            # 2. Información de municipio y departamento
            municipio, departamento = self.get_municipality_info(dataset)
            variables.update({
                'MUNICIPIO': municipio if municipio else '{MUNICIPIO}',
                'DEPARTAMENTO': departamento if departamento else '{DEPARTAMENTO}'
            })
            
            # 3. Variables de conteo base
            conteo_vars = self.get_conteo_variables(dataset, municipio if municipio else dataset)
            variables.update(conteo_vars)
            
            # 4. Variables de consistencia
            consistencia_vars = self.get_consistencia_variables(dataset)
            variables.update(consistencia_vars)
            
            # 5. Variables de omisión/comisión
            omision_vars = self.get_omision_comision_variables(dataset)
            variables.update(omision_vars)
            
            # 6. Nombre de GDB
            variables['gdb_name'] = self.get_gdb_name() or '{gdb_name}'
            
            # 7. Variables de conformidad
            variables.update({
                'consistencia_Topologica': "No Conforme" if variables.get('sum_total_err', 0) > 0 else "Conforme",
                'consistencia_formato': "No Conforme" if variables.get('total-errores', 0) > 0 else "Conforme",
                'correspondencia_geoalfa': "No Conforme" if variables.get('regis_total', 0) > 0 else "Conforme"
            })
            
            # Imprimir todas las variables para diagnóstico
            print("\nVariables procesadas:")
            for key, value in sorted(variables.items()):
                print(f"  {{{key}}}: {value}")
                
            return variables
            
        except Exception as e:
            print(f"[ERROR] Error procesando variables del dataset: {e}")
            raise

    def replace_in_word(self, dataset):
        """Reemplaza las variables en el documento Word usando la nueva estructura temporal"""
        try:
            print(f"\nProcesando documento Word para dataset {dataset}...")
            
            # Obtener variables procesadas
            variables = self.process_dataset_variables(dataset)
            
            # Buscar el documento en nuestra carpeta temporal
            docx_files = list(self.process_temp_path.glob("*.docx"))
            if not docx_files:
                raise FileNotFoundError("No se encontró ningún documento Word para procesar")
                
            template_file = docx_files[0]
            print(f"\nProcesando archivo: {template_file.name}")
            
            # Definir nombre del archivo de salida
            output_name = f"RESULTADOS_{dataset}_{template_file.name}"
            temp_output_path = self.process_temp_path / output_name
            final_output_path = self.temp_root / output_name
            
            # Cargar el documento
            doc = Document(template_file)
            replacements = 0

            def replace_text_in_paragraph(paragraph):
                """Reemplaza texto en un párrafo completo"""
                nonlocal replacements
                
                if not paragraph.runs:
                    return
                    
                # Obtener texto completo del párrafo
                full_text = paragraph.text
                new_text = full_text
                
                # Buscar todas las variables en el texto
                found_replacements = False
                for var_name, value in variables.items():
                    placeholder = f"{{{var_name}}}"
                    if placeholder in new_text:
                        try:
                            new_text = new_text.replace(placeholder, str(value))
                            replacements += 1
                            found_replacements = True
                            print(f"  [OK] Reemplazo: {placeholder} -> {value}")
                        except Exception as e:
                            print(f"  [ERROR] Error reemplazando {placeholder}: {e}")
                
                # Si se encontraron reemplazos, actualizar el párrafo
                if found_replacements:
                    try:
                        # Preservar el formato del primer run
                        first_run = paragraph.runs[0]
                        if first_run:
                            font_name = first_run.font.name
                            font_size = first_run.font.size
                            bold = first_run.font.bold
                            italic = first_run.font.italic
                            
                            # Limpiar el párrafo
                            for run in paragraph.runs[1:]:
                                run.text = ""
                            
                            # Establecer el nuevo texto con el formato preservado
                            first_run.text = new_text
                            first_run.font.name = font_name
                            first_run.font.size = font_size
                            first_run.font.bold = bold
                            first_run.font.italic = italic
                        else:
                            # Si no hay runs, crear uno nuevo
                            new_run = paragraph.add_run(new_text)
                            
                    except Exception as e:
                        print(f"  [ERROR] Error actualizando formato: {e}")

            # Procesar párrafos
            for paragraph in doc.paragraphs:
                replace_text_in_paragraph(paragraph)
            
            # Procesar tablas
            for table_index, table in enumerate(doc.tables, 1):
                for row_index, row in enumerate(table.rows, 1):
                    for cell_index, cell in enumerate(row.cells, 1):
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text
                            replace_text_in_paragraph(paragraph)
                            if original_text != paragraph.text:
                                print(f"  [OK] Reemplazo en tabla {table_index}, fila {row_index}, celda {cell_index}")

            print(f"\nTotal de reemplazos realizados: {replacements}")
            
            try:
                # Primero guardar en la carpeta temporal del proceso
                doc.save(temp_output_path)
                print(f"[OK] Documento guardado temporalmente en: {temp_output_path}")
                
                # Verificar que el archivo se guardó correctamente
                if not temp_output_path.exists():
                    raise FileNotFoundError(f"El archivo no se guardó correctamente en: {temp_output_path}")
                
                # Mover a la ubicación final en Temporary_Files
                shutil.move(str(temp_output_path), str(final_output_path))
                print(f"[OK] Documento movido a ubicación final: {final_output_path}")
                
                # Verificar que el archivo se movió correctamente
                if not final_output_path.exists():
                    raise FileNotFoundError(f"Error al mover el archivo a: {final_output_path}")
                
                # Eliminar el template original de nuestra carpeta temporal
                if template_file.exists():
                    template_file.unlink()
                    print(f"[OK] Archivo original eliminado: {template_file.name}")
                
            except Exception as e:
                print(f"[ERROR] Error guardando o moviendo el documento: {e}")
                raise
            
            return True
            
        except Exception as e:
            print(f"[ERROR] Error en replace_in_word: {e}")
            return False

    def cleanup_process_files(self):
        """Limpia los archivos de la carpeta temporal del proceso"""
        try:
            print("\nLimpiando archivos temporales del proceso...")
            
            if self.process_temp_path.exists():
                try:
                    shutil.rmtree(self.process_temp_path)
                    print(f"[OK] Carpeta temporal del proceso eliminada: {self.process_temp_path}")
                except Exception as e:
                    print(f"[ERROR] Error eliminando carpeta temporal: {e}")
                    # Intentar eliminar archivos uno por uno
                    for file in self.process_temp_path.glob("*"):
                        try:
                            file.unlink()
                            print(f"[OK] Archivo eliminado: {file.name}")
                        except Exception as e:
                            print(f"[ERROR] No se pudo eliminar {file.name}: {e}")
            
            return True
        except Exception as e:
            print(f"[ERROR] Error limpiando archivos temporales: {e}")
            return False
        

    def verify_database_integrity(self):
        """Verifica la integridad de las bases de datos para el caso urbano"""
        print("\nVerificando integridad de bases de datos...")
        
        integrity_checks = {
            'conteo': [
                "SELECT name FROM sqlite_master WHERE type='table'",
                "SELECT * FROM URBANO LIMIT 1"
            ],
            'errores': [
                "SELECT name FROM sqlite_master WHERE type='table'",
                "SELECT * FROM URBANO LIMIT 1"
            ],
            'excepciones': [
                "SELECT name FROM sqlite_master WHERE type='table'",
                "SELECT * FROM URBANO LIMIT 1"
            ],
            'omision': [
                "SELECT name FROM sqlite_master WHERE type='table'",
                "SELECT * FROM URBANO LIMIT 1"
            ],
            'municipios': [
                "SELECT name FROM sqlite_master WHERE type='table'",
                "SELECT * FROM municipios LIMIT 1"
            ]
        }
        
        results = {}
        for db_name, queries in integrity_checks.items():
            try:
                if not self.db_paths[db_name].exists():
                    results[db_name] = False
                    print(f"[X] Base de datos no encontrada: {db_name}")
                    continue
                    
                with sqlite3.connect(self.db_paths[db_name]) as conn:
                    cursor = conn.cursor()
                    for query in queries:
                        cursor.execute(query)
                results[db_name] = True
                print(f"[OK] Base de datos verificada: {db_name}")
            except Exception as e:
                results[db_name] = False
                print(f"[ERROR] Error verificando {db_name}: {e}")
        
        return all(results.values())

    def generate_diagnostic_report(self):
        """Genera un reporte de diagnóstico en la carpeta temporal"""
        try:
            report_path = self.temp_root / "word_automation_URBANO_diagnostic.txt"
            
            system_info = {
                'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'python_version': sys.version,
                'arcpy_version': arcpy.GetInstallInfo()['Version'],
                'paths': {
                    'root_path': str(self.root_path),
                    'temp_root': str(self.temp_root),
                    'process_temp': str(self.process_temp_path)
                }
            }
            
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write("=== Reporte de Diagnóstico - Proceso URBANO ===\n\n")
                f.write(f"Generado: {system_info['timestamp']}\n\n")
                
                f.write("=== Información del Sistema ===\n")
                f.write(f"Python: {system_info['python_version']}\n")
                f.write(f"ArcPy: {system_info['arcpy_version']}\n\n")
                
                f.write("=== Rutas ===\n")
                for name, path in system_info['paths'].items():
                    f.write(f"{name}: {path}\n")
                
                f.write("\n=== Bases de Datos ===\n")
                for name, path in self.db_paths.items():
                    f.write(f"{name}: {path}\n")
            
            print(f"\n[OK] Reporte de diagnóstico generado: {report_path}")
            return True
        except Exception as e:
            print(f"[ERROR] Error generando reporte de diagnóstico: {e}")
            return False

    def validate_word_document(self, doc_path):
        """Valida la estructura del documento Word"""
        try:
            doc = Document(doc_path)
            structure = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections)
            }
            
            print(f"\nEstructura del documento {doc_path.name}:")
            print(f"Párrafos: {structure['paragraphs']}")
            print(f"Tablas: {structure['tables']}")
            print(f"Secciones: {structure['sections']}")
            
            # Analizar variables en el documento
            variables_found = set()
            for paragraph in doc.paragraphs:
                matches = re.findall(r'\{([^}]+)\}', paragraph.text)
                variables_found.update(matches)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = re.findall(r'\{([^}]+)\}', paragraph.text)
                            variables_found.update(matches)
            
            print("\nVariables encontradas en el documento:")
            for var in sorted(variables_found):
                print(f"  - {{{var}}}")
            
            return True, variables_found
            
        except Exception as e:
            print(f"[ERROR] Error validando documento Word: {e}")
            return False, set()


    def get_rule_mapping(self):
        """
        Retorna un diccionario con el mapeo de reglas inglés/español ORIGINAL
        """
        return {
            'Must Not Have Gaps': 'No debe tener espacios',
            'Must Not Overlap': 'No debe superponerse',
            'Must Be Covered By Feature Class Of': 'Debe ser cubierto por la clase de entidad de',
            'Must Cover Each Other': 'Deben cubrirse entre ellos'
        }

    def build_rule_condition(self, rule_name):
        """
        Construye la condición SQL para una regla específica considerando ambos idiomas
        """
        rule_mapping = self.get_rule_mapping()
        english_rule = rule_name
        spanish_rule = rule_mapping.get(rule_name, '')
        # Notar que mantenemos ambas versiones de la columna (ruledescription/RuleDescription)
        return f"""(ruledescription = '{english_rule}' 
                OR ruledescription = '{spanish_rule}'
                OR RuleDescription = '{english_rule}'
                OR RuleDescription = '{spanish_rule}')"""

    def get_omision_comision_variables(self, dataset):
        """Obtiene variables de omisión/comisión con manejo mejorado de errores para URBANO"""
        print(f"\nObteniendo variables de omisión/comisión para {dataset}")
        results = {}
        
        try:
            # Definir las consultas con sus descripciones para URBANO
            queries = {
                'regis1': {
                    'query': "SELECT SUM(omision_terrenos) AS total_omision_terrenos FROM URBANO",
                    'description': "Omisión de Terrenos"
                },
                'regis2': {
                    'query': "SELECT SUM(comision_terrenos) AS total_omision_terrenos FROM URBANO",
                    'description': "Comisión de Terrenos"
                },
                'regis3': {
                    'query': "SELECT SUM(omision_unidades) AS total_omision_terrenos FROM URBANO",
                    'description': "Omisión de Unidades de Construcción"
                },
                'regis4': {
                    'query': "SELECT SUM(comision_unidades) AS total_omision_terrenos FROM URBANO",
                    'description': "Comisión de Unidades de Construcción"
                },
                'regis5': {
                    'query': "SELECT SUM(duplicados) AS total_omision_terrenos FROM URBANO",
                    'description': "Elementos Duplicados"
                }
            }
            
            # Procesar cada consulta
            for var_name, query_info in queries.items():
                try:
                    value = self.execute_sql_query(self.db_paths['omision'], query_info['query'])
                    results[var_name] = value
                    print(f"[OK] {var_name} ({query_info['description']}): {value}")
                except Exception as e:
                    print(f"[ERROR] Error procesando {var_name} ({query_info['description']}): {e}")
                    results[var_name] = 0
            
            # Calcular total con validación
            try:
                results['regis_total'] = sum(results[f'regis{i}'] for i in range(1, 6))
                print(f"Total registros calculado: {results['regis_total']}")
            except Exception as e:
                print(f"[ERROR] Error calculando total de registros: {e}")
                results['regis_total'] = 0
            
            return results
            
        except Exception as e:
            print(f"[ERROR] Error general en get_omision_comision_variables: {e}")
            return {'regis1': 0, 'regis2': 0, 'regis3': 0, 'regis4': 0, 'regis5': 0, 'regis_total': 0}
        

    def get_error_queries():
        """
        Genera las consultas para errores con soporte bilingüe
        """
        return {
            'no_huecos': f"""
                SELECT COUNT(*) AS cantidad 
                FROM URBANO 
                WHERE {build_rule_condition('Must Not Have Gaps')}
            """,
            'no_superponer': f"""
                SELECT COUNT(*) AS cantidad 
                FROM URBANO 
                WHERE {build_rule_condition('Must Not Overlap')}
            """,
            'cubierto_por': f"""
                SELECT COUNT(*) AS cantidad 
                FROM URBANO 
                WHERE {build_rule_condition('Must Be Covered By Feature Class Of')}
            """,
            'cubrirse_entre': f"""
                SELECT COUNT(*) AS cantidad 
                FROM URBANO 
                WHERE {build_rule_condition('Must Cover Each Other')}
            """
        }

    def get_exception_queries():
        """
        Genera las consultas para excepciones con soporte bilingüe
        """
        return {
            'no_huecos_': f"""
                SELECT COUNT(*) 
                FROM "URBANO" 
                WHERE ({build_rule_condition('Must Not Have Gaps')})
                AND "isException" != 0
            """,
            'no_superponer_': f"""
                SELECT COUNT(*) 
                FROM "URBANO" 
                WHERE ({build_rule_condition('Must Not Overlap')})
                AND "isException" != 0
            """,
            'cubierto_por_': f"""
                SELECT COUNT(*) 
                FROM "URBANO" 
                WHERE ({build_rule_condition('Must Be Covered By Feature Class Of')})
                AND "isException" != 0
            """,
            'cubrirse_entre_': f"""
                SELECT COUNT(*) 
                FROM "URBANO" 
                WHERE ({build_rule_condition('Must Cover Each Other')})
                AND "isException" != 0
            """
        }

    # Ejemplo de uso en tu código existente:
    def process_topology_errors(db_connection):
        """
        Procesa los errores topológicos considerando ambos idiomas
        """
        try:
            error_queries = get_error_queries()
            exception_queries = get_exception_queries()
            results = {}
            
            # Procesar errores
            for query_name, query in error_queries.items():
                try:
                    cursor = db_connection.cursor()
                    cursor.execute(query)
                    results[query_name] = cursor.fetchone()[0]
                    print(f"[OK] {query_name}: {results[query_name]}")
                except Exception as e:
                    print(f"[ERROR] Error procesando {query_name}: {e}")
                    results[query_name] = 0

            # Procesar excepciones
            for query_name, query in exception_queries.items():
                try:
                    cursor = db_connection.cursor()
                    cursor.execute(query)
                    results[query_name] = cursor.fetchone()[0]
                    print(f"[OK] {query_name}: {results[query_name]}")
                except Exception as e:
                    print(f"[ERROR] Error procesando {query_name}: {e}")
                    results[query_name] = 0

            return results

        except Exception as e:
            print(f"[ERROR] Error general procesando errores topológicos: {e}")
            return {}


    def run(self):
        """Método principal de ejecución"""
        try:
            print(f"\n{'='*50}")
            print("Iniciando proceso de automatización URBANO...")
            
            # Generar reporte de diagnóstico inicial
            self.generate_diagnostic_report()
            
            # Verificar rutas
            if not self.verify_paths():
                print("\n[ADVERTENCIA] Algunas rutas necesarias no existen.")
                return
            
            # Verificar integridad de bases de datos
            if not self.verify_database_integrity():
                print("\n[ERROR] Verificación de integridad de bases de datos fallida")
                return
            
            # Copiar plantillas
            self.copy_templates()
            
            # Verificar si se copiaron las plantillas
            docx_files = list(self.process_temp_path.glob("*.docx"))
            if not docx_files:
                print("\n[ADVERTENCIA] No se encontraron documentos Word para procesar.")
                return
            
            print(f"\nArchivos Word encontrados para procesar:")
            for doc in docx_files:
                print(f"  - {doc.name}")
            
            # Procesar cada dataset
            for dataset in self.datasets_to_process:
                try:
                    print(f"\nProcesando dataset: {dataset}")
                    print("-" * 30)
                    
                    # Validar el documento template antes de procesarlo
                    template = next(iter(docx_files), None)
                    if template:
                        is_valid, variables = self.validate_word_document(template)
                        if not is_valid:
                            print(f"[ADVERTENCIA] El documento template puede tener problemas")
                    
                    if self.replace_in_word(dataset):
                        print(f"[OK] Procesamiento de {dataset} completado")
                    else:
                        print(f"[ERROR] Error procesando {dataset}")
                        
                except Exception as e:
                    print(f"[ERROR] Error procesando dataset {dataset}: {e}")
                    print(traceback.format_exc())
            
            # Limpiar archivos temporales del proceso
            self.cleanup_process_files()
            
            print(f"\n{'='*50}")
            print("Proceso URBANO Finalizado")
            print(f"{'='*50}")
            
        except Exception as e:
            print(f"\n[ERROR] Error durante la ejecución: {e}")
            print(traceback.format_exc())
            raise

# Código principal
if __name__ == "__main__":
    try:
        # Configurar manejo de excepciones no capturadas
        def handle_exception(exc_type, exc_value, exc_traceback):
            print(f"\n[ERROR CRÍTICO NO CAPTURADO] {exc_value}")
            print("".join(traceback.format_tb(exc_traceback)))
        
        sys.excepthook = handle_exception
        
        # Iniciar el proceso
        automation = WordAutomation()
        automation.run()
        
    except Exception as e:
        print(f"\n[ERROR CRÍTICO] {e}")
        print(traceback.format_exc())
        
    finally:
        print("Informe Listo para Exportar")
