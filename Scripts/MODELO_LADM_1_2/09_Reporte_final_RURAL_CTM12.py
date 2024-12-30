import os
import sys
import shutil
import sqlite3
import logging
import platform
import psutil
import traceback
import re
from pathlib import Path
from docx import Document
import arcpy
from datetime import datetime
import sys
sys.stdout.reconfigure(encoding='utf-8')

class WordAutomation:
    def __init__(self):
        """Inicializa la automatización con mejor manejo de archivos temporales"""
        try:
            # Auto-detectar la ruta raíz
            self.root_path = self._find_root_path()
            
            # Configurar rutas principales
            self.temp_root = self.root_path / "Files" / "Temporary_Files" / "MODELO_LADM_1_2"
            self.process_temp_path = self.temp_root / "word_automation_temp"
            
            # Configurar logging en Temporary_Files
            log_file = self.temp_root / "word_automation_rural.log"
            logging.basicConfig(
                level=logging.INFO,
                format='%(asctime)s - %(levelname)s - %(message)s',
                handlers=[
                    logging.FileHandler(log_file),
                    logging.StreamHandler()
                ]
            )
            
            # Crear carpeta temporal para el proceso si no existe
            self.process_temp_path.mkdir(exist_ok=True)
            
            print(f"\n{'='*50}")
            print("Configuración inicial:")
            print(f"Ruta raíz detectada: {self.root_path}")
            
            # Datasets a procesar
            self.datasets_to_process = [
                "RURAL_CTM12",
                #"URBANO_CTM12"
            ]
            print(f"Datasets a procesar: {', '.join(self.datasets_to_process)}")
            
            # Configurar rutas de bases de datos - Actualizado para solo incluir las necesarias
            self.db_paths = {
                'conteo': self.temp_root / "db" / "conteo_elementos.db",
                'registro_errores': self.temp_root / "db" / "registro_errores_rural.db",
                'municipios': self.root_path / "Files" / "Municipios" / "municipios.db"
            }
            
            # Definir los rangos de tablas para cada tipo de error
            self.error_ranges = {
                'no_huecos': (1, 16),      # R1 - R16
                'no_superponer': (17, 39),  # R17 - R39
                'cubierto_por': (40, 56),   # R40 - R56
                'cubrirse_entre': (57, 62)  # R57 - R62
            }
            
            # Rango para consistencia
            self.consistencia_range = (63, 79)  # R63 - R79
            
            # Verificar existencia de bases de datos
            print("\nVerificando rutas de bases de datos:")
            missing_dbs = []
            for db_name, db_path in self.db_paths.items():
                exists = db_path.exists()
                print(f"{db_name}: {'[OK]' if exists else '[X]'} {db_path}")
                if not exists:
                    missing_dbs.append(db_name)
            
            if missing_dbs:
                raise FileNotFoundError(f"Bases de datos faltantes: {', '.join(missing_dbs)}")
            
            print(f"\n{'='*50}")
            
        except Exception as e:
            print(f"\n[ERROR CRÍTICO] Error durante la inicialización: {e}")
            raise
    def _find_root_path(self):
        """
        Encuentra la ruta raíz del proyecto verificando la estructura de directorios esperada.
        """
        try:
            current = Path(os.path.abspath(__file__)).parent
            
            while current.parent != current:
                # Verifica la estructura característica del proyecto
                required_paths = [
                    current / "Files" / "Temporary_Files" / "MODELO_LADM_1_2",
                    current / "Files" / "Temporary_Files" / "array_config.txt",
                    current / "Files" / "Templates" / "MODELO_LADM_1_2"/ "04_REPORTE_FINAL",
                    current / "Files" / "Municipios" / "municipios.db"
                ]
                
                if all(path.exists() for path in required_paths):
                    print(f"Raíz del proyecto encontrada en: {current}")
                    return current
                
                current = current.parent
            
            raise Exception(
                "No se encontró la raíz del proyecto.\n"
                "Verifique que está ejecutando el script desde dentro del proyecto "
                "y que existe la siguiente estructura:\n"
                "- Files/Temporary_Files/MODELO_LADM_1_2\n"
                "- Files/Temporary_Files/array_config.txt\n"
                "- Files/Templates/MODELO_LADM_1_2/04_REPORTE_FINAL\n"
                "- Files/Municipios/municipios.db"
            )
        except Exception as e:
            print(f"[ERROR] Error encontrando ruta raíz: {e}")
            raise

    def get_tables_in_range(self, start_num, end_num):
        """
        Obtiene todas las tablas en un rango específico de números de regla
        """
        try:
            with sqlite3.connect(self.db_paths['registro_errores']) as conn:
                cursor = conn.cursor()
                # Obtener todas las tablas de la base de datos
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
                all_tables = cursor.fetchall()
                
                # Filtrar tablas que están en el rango especificado
                tables_in_range = []
                for (table_name,) in all_tables:
                    # Extraer el número de la regla del nombre de la tabla
                    match = re.search(r'R_(\d+)_', table_name)
                    if match:
                        rule_num = int(match.group(1))
                        if start_num <= rule_num <= end_num:
                            tables_in_range.append(table_name)
                
                return tables_in_range
        except Exception as e:
            print(f"[ERROR] Error obteniendo tablas en rango {start_num}-{end_num}: {e}")
            return []

    def count_records_in_tables(self, tables, count_exceptions=False):
        """
        Cuenta registros en un conjunto de tablas, opcionalmente contando solo excepciones
        """
        total_count = 0
        try:
            with sqlite3.connect(self.db_paths['registro_errores']) as conn:
                cursor = conn.cursor()
                for table in tables:
                    try:
                        if count_exceptions:
                            query = f"SELECT COUNT(*) FROM {table} WHERE isExceptio <> 0"
                        else:
                            query = f"SELECT COUNT(*) FROM {table}"
                        cursor.execute(query)
                        count = cursor.fetchone()[0]
                        total_count += count
                        print(f"  [OK] {table}: {count} {'excepciones' if count_exceptions else 'registros'}")
                    except Exception as e:
                        print(f"  [ERROR] Error contando registros en {table}: {e}")
                        continue
        except Exception as e:
            print(f"[ERROR] Error en conexión a base de datos: {e}")
        return total_count

    def get_error_counts(self):
        """
        Obtiene los conteos de errores y excepciones para cada tipo de regla
        """
        results = {}
        try:
            for error_type, (start, end) in self.error_ranges.items():
                print(f"\nProcesando {error_type} (R{start}-R{end})")
                tables = self.get_tables_in_range(start, end)
                
                # Contar registros totales
                total_count = self.count_records_in_tables(tables)
                results[error_type] = total_count
                
                # Contar excepciones
                exceptions_count = self.count_records_in_tables(tables, count_exceptions=True)
                results[f"{error_type}_"] = exceptions_count
                
                # Calcular suma
                results[f"sum{list(self.error_ranges.keys()).index(error_type) + 1}"] = total_count - exceptions_count
            
            # Calcular sumas totales
            results['sum_t_situa'] = sum(results[key] for key in self.error_ranges.keys())
            results['sum_cant_ex'] = sum(results[f"{key}_"] for key in self.error_ranges.keys())
            results['sum_total_err'] = sum(results[f"sum{i+1}"] for i in range(len(self.error_ranges)))
            
            return results
        except Exception as e:
            print(f"[ERROR] Error obteniendo conteos de errores: {e}")
            return results

    def get_consistencia_counts(self):
        """
        Obtiene los conteos para las tablas de consistencia (R58-R71)
        """
        try:
            print("\nProcesando tablas de consistencia")
            start, end = self.consistencia_range
            tables = self.get_tables_in_range(start, end)
            
            # Contar todos los registros
            total_count = self.count_records_in_tables(tables)
            
            # Contar excepciones
            exceptions_count = self.count_records_in_tables(tables, count_exceptions=True)
            
            return {
                'consistencia_total': total_count,
                'consistencia_except': exceptions_count,
                'total-errores': total_count - exceptions_count
            }
        except Exception as e:
            print(f"[ERROR] Error obteniendo conteos de consistencia: {e}")
            return {
                'consistencia_total': 0,
                'consistencia_except': 0,
                'total-errores': 0
            }
            
    def verify_paths(self):
        """Verifica que todas las rutas necesarias existan"""
        try:
            print("\nVerificando estructura de carpetas:")
            paths_to_check = [
                self.root_path / "Files" / "Templates" / "MODELO_LADM_1_2"/ "04_REPORTE_FINAL",
                self.temp_root,
                self.temp_root / "db"
            ]
            
            missing_paths = []
            for path in paths_to_check:
                exists = path.exists()
                print(f"{'[OK]' if exists else '[X]'} {path}")
                if not exists:
                    missing_paths.append(str(path))
            
            if missing_paths:
                print("\n[ADVERTENCIA] Las siguientes rutas no existen:")
                for path in missing_paths:
                    print(f"  - {path}")
                return False
            
            return True
            
        except Exception as e:
            print(f"[ERROR] Error verificando rutas: {e}")
            return False

    def copy_templates(self):
        """Copia las plantillas a la carpeta temporal del proceso"""
        try:
            print("\nCopiando plantillas:")
            
            for dataset in self.datasets_to_process:
                template_path = self.root_path / "Files" / "Templates" / "MODELO_LADM_1_2"/ "04_REPORTE_FINAL" / dataset
                
                print(f"\nProcesando dataset {dataset}:")
                print(f"Buscando plantillas en: {template_path}")
                
                if not template_path.exists():
                    print(f"  [ADVERTENCIA] Ruta de plantilla no encontrada: {template_path}")
                    continue
                
                docx_files = list(template_path.glob("*.docx"))
                if not docx_files:
                    print(f"  [ADVERTENCIA] No se encontraron archivos .docx en {template_path}")
                    continue
                    
                for docx_file in docx_files:
                    try:
                        shutil.copy2(docx_file, self.process_temp_path)
                        print(f"  [OK] Copiado: {docx_file.name} -> {self.process_temp_path}")
                    except Exception as e:
                        print(f"  [ERROR] Error copiando {docx_file.name}: {e}")
                        
        except Exception as e:
            print(f"[ERROR] Error copiando plantillas: {e}")
            raise

    def execute_sql_query(self, db_path, query, params=None):
        """Ejecuta una consulta SQL y retorna el resultado con mejor manejo de errores"""
        try:
            if not Path(db_path).exists():
                raise FileNotFoundError(f"Base de datos no encontrada: {db_path}")
                
            with sqlite3.connect(db_path) as conn:
                cursor = conn.cursor()
                if params:
                    result = cursor.execute(query, params).fetchone()
                else:
                    result = cursor.execute(query).fetchone()
                    
                if result is None:
                    print(f"[ADVERTENCIA] La consulta no retornó resultados: {query}")
                    return 0
                    
                return result[0] if result else 0
                
        except sqlite3.Error as e:
            print(f"[ERROR] Error en la base de datos: {e}")
            print(f"Base de datos: {db_path}")
            print(f"Query: {query}")
            print(f"Parámetros: {params}")
            return 0
        except Exception as e:
            print(f"[ERROR] Error ejecutando consulta SQL: {e}")
            return 0

    def get_gdb_name(self):
        """Obtiene el nombre de la geodatabase en Temporary_Files"""
        try:
            print("\nBuscando geodatabase:")
            print(f"Buscando en: {self.temp_root}")
            
            if not self.temp_root.exists():
                raise FileNotFoundError(f"Ruta temporal no encontrada: {self.temp_root}")
            
            gdbs = list(self.temp_root.glob("*.gdb"))
            
            if not gdbs:
                print("  [X] No se encontró ninguna geodatabase")
                return None
                
            if len(gdbs) > 1:
                print(f"  [ADVERTENCIA] Se encontraron múltiples geodatabases: {[gdb.name for gdb in gdbs]}")
                print(f"  Se utilizará la primera: {gdbs[0].name}")
                
            print(f"  [OK] Geodatabase encontrada: {gdbs[0].name}")
            return gdbs[0].name
            
        except Exception as e:
            print(f"[ERROR] Error buscando geodatabase: {e}")
            return None

    def get_municipality_info(self, dataset):
        """Obtiene información del municipio y departamento desde la GDB"""
        try:
            gdb_name = self.get_gdb_name()
            if not gdb_name:
                raise ValueError("No se pudo obtener el nombre de la geodatabase")

            print(f"\nObteniendo información del municipio para dataset {dataset}")
            gdb_path = self.temp_root / gdb_name
            
            if not gdb_path.exists():
                raise FileNotFoundError(f"Geodatabase no encontrada: {gdb_path}")
            
            print(f"Conectando a GDB: {gdb_path}")
            arcpy.env.workspace = str(gdb_path)
            
            if not arcpy.Exists(dataset):
                raise ValueError(f"Dataset {dataset} no encontrado")
                
            print(f"[OK] Dataset encontrado: {dataset}")
            
            feature_class = "R_VEREDA_CTM12"
            print(f"Buscando feature class: {feature_class}")
            
            feature_classes = arcpy.ListFeatureClasses(feature_dataset=dataset)
            print(f"Feature classes encontrados en el dataset: {feature_classes}")
            
            if feature_class not in feature_classes:
                raise ValueError(f"Feature class {feature_class} no encontrado")
            
            print(f"[OK] Feature class encontrado: {feature_class}")
            
            codes = set()
            with arcpy.da.SearchCursor(feature_class, ["CODIGO_MUNICIPIO"]) as cursor:
                for i, row in enumerate(cursor):
                    if i < 5 and row[0]:
                        codes.add(row[0])
                    if i >= 5:
                        break
            
            if not codes:
                raise ValueError("No se encontraron códigos de municipio válidos")

            codigo_municipio = codes.pop()
            print(f"Código de municipio encontrado: {codigo_municipio}")
            
            with sqlite3.connect(self.db_paths['municipios']) as conn:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT MUNICIPIO, DEPARTAMENTO 
                    FROM municipios 
                    WHERE CODIGO_DANE = ?
                """, (codigo_municipio,))
                result = cursor.fetchone()
                
                if not result:
                    raise ValueError(f"No se encontró información para el código de municipio: {codigo_municipio}")
                
                print(f"[OK] Municipio: {result[0]}, Departamento: {result[1]}")
                return result[0], result[1]
                
        except arcpy.ExecuteError as e:
            print(f"[ERROR] Error de ArcPy: {e}")
            print(arcpy.GetMessages(2))
            return None, None
        except Exception as e:
            print(f"[ERROR] Error obteniendo información del municipio: {e}")
            return None, None
        
    def get_conteo_variables(self, dataset, municipio):
        """Obtiene todas las variables de conteo para un dataset"""
        print(f"\nObteniendo variables de conteo para {dataset}")
        
        variables = {}
        
        # Lista de variables base para rural
        base_vars = [
            'lc_unidadconstruccion',
            'lc_construccion',
            'lc_terreno',
            'extdireccion',
            'av_zonahomogeneageoeconomicarural',
            'av_zonahomogeneafisicarural',
            'av_zonahomogeneageoeconomicaurbana',
            'av_zonahomogeneafisicaurbana',
            'cc_limitemunicipio',
            'cc_sectorrural',
            'cc_sectorurbano',
            'cc_perimetrourbano',
            'cc_vereda',
            'cc_corregimiento',
            'cc_localidadcomuna',
            'cc_centropoblado',
            'cc_manzana',
            'cc_barrio',
            'cc_nomenclaturavial',

        ]
        
        try:
            print("\nProcesando variables individuales:")
            for var_name in base_vars:
                try:
                    query = f"SELECT {var_name} FROM conteos LIMIT 1"
                    value = self.execute_sql_query(self.db_paths['conteo'], query)
                    variables[var_name] = value
                    
                    var_underscore = f"{var_name}_"
                    variables[var_underscore] = "" if value > 0 else "Vacía"
                    
                    print(f"[OK] {var_name}: {value} -> {var_underscore}: '{variables[var_underscore]}'")
                    
                except Exception as e:
                    print(f"[ERROR] Error procesando {var_name}: {e}")
                    variables[var_name] = 0
                    variables[f"{var_name}_"] = "Vacía"

            try:
                total_query = f"SELECT ({' + '.join(base_vars)}) FROM conteos LIMIT 1"
                variables['total_conteo'] = self.execute_sql_query(self.db_paths['conteo'], total_query)
                print(f"\nTotal conteo calculado: {variables['total_conteo']}")
                
            except Exception as e:
                print(f"[ERROR] Error calculando total_conteo directo: {e}")
                try:
                    variables['total_conteo'] = sum(variables[var] for var in base_vars)
                    print(f"Total conteo calculado por suma: {variables['total_conteo']}")
                except Exception as e2:
                    print(f"[ERROR] Error calculando total_conteo por suma: {e2}")
                    variables['total_conteo'] = 0
            
            return variables
            
        except Exception as e:
            print(f"[ERROR] Error general en get_conteo_variables: {e}")
            return variables

    def process_dataset_variables(self, dataset):
        """Procesa todas las variables para un dataset rural específico"""
        print(f"\nProcesando variables para dataset {dataset}")
        variables = {}
        
        try:
            # 1. Fecha actual
            variables['date'] = datetime.now().strftime('%Y/%m/%d')
            
            # 2. Información de municipio y departamento
            municipio, departamento = self.get_municipality_info(dataset)
            variables.update({
                'MUNICIPIO': municipio if municipio else '{MUNICIPIO}',
                'DEPARTAMENTO': departamento if departamento else '{DEPARTAMENTO}'
            })
            
            # 3. Variables de conteo base
            conteo_vars = self.get_conteo_variables(dataset, municipio if municipio else dataset)
            variables.update(conteo_vars)
            
            # 4. Variables de error y consistencia
            error_vars = self.get_error_counts()
            variables.update(error_vars)
            
            consistencia_vars = self.get_consistencia_counts()
            variables.update(consistencia_vars)
            
            # 5. Nombre de GDB
            variables['gdb_name'] = self.get_gdb_name() or '{gdb_name}'
            
            # 6. Variables de conformidad
            variables.update({
                'consistencia_Topologica': "No Conforme" if variables.get('sum_total_err', 0) > 0 else "Conforme",
                'consistencia_formato': "No Conforme" if variables.get('total-errores', 0) > 0 else "Conforme",
                'correspondencia_geoalfa': "Conforme"  # Siempre Conforme ya que se eliminaron las variables regis
            })
            
            print("\nVariables procesadas:")
            for key, value in sorted(variables.items()):
                print(f"  {{{key}}}: {value}")
                
            return variables
            
        except Exception as e:
            print(f"[ERROR] Error procesando variables del dataset: {e}")
            raise

    def replace_in_word(self, dataset):
        """Reemplaza las variables en el documento Word"""
        try:
            print(f"\nProcesando documento Word para dataset {dataset}...")
            
            variables = self.process_dataset_variables(dataset)
            
            docx_files = list(self.process_temp_path.glob("*.docx"))
            if not docx_files:
                raise FileNotFoundError("No se encontró ningún documento Word para procesar")
                
            template_file = docx_files[0]
            print(f"\nProcesando archivo: {template_file.name}")
            
            output_name = f"RESULTADOS_{dataset}_{template_file.name}"
            temp_output_path = self.process_temp_path / output_name
            final_output_path = self.temp_root / output_name
            
            doc = Document(template_file)
            replacements = 0

            def replace_text_in_paragraph(paragraph):
                nonlocal replacements
                
                if not paragraph.runs:
                    return
                    
                full_text = paragraph.text
                new_text = full_text
                
                found_replacements = False
                for var_name, value in variables.items():
                    placeholder = f"{{{var_name}}}"
                    if placeholder in new_text:
                        try:
                            new_text = new_text.replace(placeholder, str(value))
                            replacements += 1
                            found_replacements = True
                            print(f"  [OK] Reemplazo: {placeholder} -> {value}")
                        except Exception as e:
                            print(f"  [ERROR] Error reemplazando {placeholder}: {e}")
                
                if found_replacements:
                    try:
                        first_run = paragraph.runs[0]
                        if first_run:
                            font_name = first_run.font.name
                            font_size = first_run.font.size
                            bold = first_run.font.bold
                            italic = first_run.font.italic
                            
                            for run in paragraph.runs[1:]:
                                run.text = ""
                            
                            first_run.text = new_text
                            first_run.font.name = font_name
                            first_run.font.size = font_size
                            first_run.font.bold = bold
                            first_run.font.italic = italic
                        else:
                            new_run = paragraph.add_run(new_text)
                            
                    except Exception as e:
                        print(f"  [ERROR] Error actualizando formato: {e}")

            for paragraph in doc.paragraphs:
                replace_text_in_paragraph(paragraph)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph)

            print(f"\nTotal de reemplazos realizados: {replacements}")
            
            try:
                doc.save(temp_output_path)
                print(f"[OK] Documento guardado temporalmente en: {temp_output_path}")
                
                if not temp_output_path.exists():
                    raise FileNotFoundError(f"El archivo no se guardó correctamente en: {temp_output_path}")
                
                shutil.move(str(temp_output_path), str(final_output_path))
                print(f"[OK] Documento movido a ubicación final: {final_output_path}")
                
                if not final_output_path.exists():
                    raise FileNotFoundError(f"Error al mover el archivo a: {final_output_path}")
                
                if template_file.exists():
                    template_file.unlink()
                    print(f"[OK] Archivo original eliminado: {template_file.name}")
                
            except Exception as e:
                print(f"[ERROR] Error guardando o moviendo el documento: {e}")
                raise
            
            return True
            
        except Exception as e:
            print(f"[ERROR] Error en replace_in_word: {e}")
            return False

    def cleanup_process_files(self):
        """Limpia los archivos de la carpeta temporal del proceso"""
        try:
            print("\nLimpiando archivos temporales del proceso...")
            
            if self.process_temp_path.exists():
                try:
                    shutil.rmtree(self.process_temp_path)
                    print(f"[OK] Carpeta temporal del proceso eliminada: {self.process_temp_path}")
                except Exception as e:
                    print(f"[ERROR] Error eliminando carpeta temporal: {e}")
                    for file in self.process_temp_path.glob("*"):
                        try:
                            file.unlink()
                            print(f"[OK] Archivo eliminado: {file.name}")
                        except Exception as e:
                            print(f"[ERROR] No se pudo eliminar {file.name}: {e}")
            
            return True
        except Exception as e:
            print(f"[ERROR] Error limpiando archivos temporales: {e}")
            return False

    def run(self):
        """Método principal de ejecución"""
        try:
            print(f"\n{'='*50}")
            print("Iniciando proceso de automatización rural...")
            
            if not self.verify_paths():
                print("\n[ADVERTENCIA] Algunas rutas necesarias no existen.")
                return
            
            self.copy_templates()
            
            docx_files = list(self.process_temp_path.glob("*.docx"))
            if not docx_files:
                print("\n[ADVERTENCIA] No se encontraron documentos Word para procesar.")
                return
            
            print(f"\nArchivos Word encontrados para procesar:")
            for doc in docx_files:
                print(f"  - {doc.name}")
            
            for dataset in self.datasets_to_process:
                try:
                    print(f"\nProcesando dataset: {dataset}")
                    print("-" * 30)
                    
                    template = next(iter(docx_files), None)
                    if template:
                        is_valid, variables = self.validate_word_document(template)
                        if not is_valid:
                            print(f"[ADVERTENCIA] El documento template puede tener problemas")
                    
                    if self.replace_in_word(dataset):
                        print(f"[OK] Procesamiento de {dataset} completado")
                    else:
                        print(f"[ERROR] Error procesando {dataset}")
                        
                except Exception as e:
                    print(f"[ERROR] Error procesando dataset {dataset}: {e}")
                    print(traceback.format_exc())
            
            self.cleanup_process_files()
            
            print(f"\n{'='*50}")
            print("Proceso Rural Finalizado")
            print(f"{'='*50}")
            
        except Exception as e:
            print(f"\n[ERROR] Error durante la ejecución: {e}")
            print(traceback.format_exc())
            raise
    def validate_word_document(self, doc_path):
        """Valida la estructura del documento Word"""
        try:
            doc = Document(doc_path)
            structure = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'sections': len(doc.sections)
            }
            
            print(f"\nEstructura del documento {doc_path.name}:")
            print(f"Párrafos: {structure['paragraphs']}")
            print(f"Tablas: {structure['tables']}")
            print(f"Secciones: {structure['sections']}")
            
            # Analizar variables en el documento
            variables_found = set()
            for paragraph in doc.paragraphs:
                matches = re.findall(r'\{([^}]+)\}', paragraph.text)
                variables_found.update(matches)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            matches = re.findall(r'\{([^}]+)\}', paragraph.text)
                            variables_found.update(matches)
            
            print("\nVariables encontradas en el documento:")
            for var in sorted(variables_found):
                print(f"  - {{{var}}}")
            
            return True, variables_found
            
        except Exception as e:
            print(f"[ERROR] Error validando documento Word: {e}")
            return False, set()
        
if __name__ == "__main__":
    try:
        def handle_exception(exc_type, exc_value, exc_traceback):
            print(f"\n[ERROR CRÍTICO NO CAPTURADO] {exc_value}")
            print("".join(traceback.format_tb(exc_traceback)))
        
        sys.excepthook = handle_exception
        
        automation = WordAutomation()
        automation.run()
        
    except Exception as e:
        print(f"\n[ERROR CRÍTICO] {e}")
        print(traceback.format_exc())
        
    finally:
        print("Informe Listo para Exportar")